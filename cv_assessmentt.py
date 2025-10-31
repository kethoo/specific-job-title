import os
import json
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
import PyPDF2
import openai

from models import CandidateAssessment


class CVAssessmentSystem:
    def __init__(self, api_key: str = None):
        """Initialize the CV assessment system"""
        self.client = openai.OpenAI(api_key=api_key or os.getenv("OPENAI_API_KEY"))
        self.job_requirements = ""
        self.assessments: List[Any] = []
        self.session_id = f"assessment_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        self.custom_criteria = None  # Store custom criteria

    # ------------------- LOAD JOB REQUIREMENTS -------------------

    def load_job_requirements(self, file_path: str) -> str:
        """Load tender or job requirements from Word or PDF file (auto-detects type)."""
        with open(file_path, "rb") as f:
            header = f.read(4)
        try:
            if header.startswith(b"%PDF"):
                text = self._extract_text_from_pdf(file_path)
            else:
                text = self._extract_text_from_word(file_path)
        except Exception as e:
            raise ValueError(f"Cannot read file {file_path}: {e}")

        self.job_requirements = text
        print(f"✅ Loaded job requirements from: {file_path}")
        return text

    # ------------------- PROCESS CV FOLDER -------------------

    def process_cv_folder(self, folder_path: str, mode: str = "structured", custom_criteria: Dict = None) -> List[Any]:
        """Process all CVs in a folder"""
        self.custom_criteria = custom_criteria  # Store for use in assessment
        
        cv_files = []
        for ext in ["*.pdf", "*.PDF", "*.doc", "*.DOC", "*.docx", "*.DOCX"]:
            cv_files.extend(Path(folder_path).glob(ext))

        print(f"🔍 Found {len(cv_files)} CV files in {folder_path}")
        if not cv_files:
            print("⚠️ No CV files found! Check uploads.")
            return []

        for cv_file in cv_files:
            try:
                print(f"🧾 Processing: {cv_file.name}")
                cv_text = self._extract_cv_text(cv_file)
                if mode == "critical":
                    report = self._assess_candidate_critical(cv_file.name, cv_text)
                    self.assessments.append(report)
                else:
                    assessment = self._assess_candidate_structured(cv_file.name, cv_text)
                    self.assessments.append(assessment)
            except Exception as e:
                print(f"❌ Error processing {cv_file.name}: {e}")
                continue

        print(f"✅ Completed assessments for {len(self.assessments)} candidates.")
        return self.assessments

    # ------------------- FILE HELPERS -------------------

    def _extract_cv_text(self, file_path: Path) -> str:
        """Extract text from CV file"""
        if file_path.suffix.lower() == ".pdf":
            return self._extract_text_from_pdf(str(file_path))
        return self._extract_text_from_word(str(file_path))

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = []
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)

    def _extract_text_from_word(self, file_path: str) -> str:
        """Extract full text from a Word (.docx) file including tables, text boxes, and nested items."""
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        import itertools
    
        doc = Document(file_path)
    
        def iter_block_items(parent):
            """Yield paragraphs and tables in document order (recursively for nested tables)."""
            for child in parent.element.body.iterchildren():
                if isinstance(child, CT_P):
                    yield parent.paragraphs[len(list(itertools.takewhile(lambda p: p._p is not child, parent.paragraphs)))]
                elif isinstance(child, CT_Tbl):
                    yield parent.tables[len(list(itertools.takewhile(lambda t: t._tbl is not child, parent.tables)))]
    
        def get_text_from_cell(cell):
            parts = []
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text.strip())
            for table in cell.tables:
                parts.append(get_text_from_table(table))
            return " ".join(parts)
    
        def get_text_from_table(table):
            rows_text = []
            for row in table.rows:
                row_text = " ".join(get_text_from_cell(c) for c in row.cells if c.text.strip())
                if row_text:
                    rows_text.append(row_text)
            return " ".join(rows_text)
    
        text_blocks = []
        for block in iter_block_items(doc):
            if isinstance(block, type(doc.paragraphs[0])):
                if block.text.strip():
                    text_blocks.append(block.text.strip())
            elif isinstance(block, type(doc.tables[0])):
                table_text = get_text_from_table(block)
                if table_text:
                    text_blocks.append(table_text)
    
        # Clean and merge
        text = " ".join(text_blocks)
        text = re.sub(r"\s{2,}", " ", text)
        text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)
        return text.strip()
    
        
        


    # ------------------- STRUCTURED (DASHBOARD) MODE -------------------

    def _assess_candidate_structured(self, filename: str, cv_text: str) -> CandidateAssessment:
        """Structured scoring (dashboard mode)"""
        prompt = f"""
You are an HR evaluator performing a structured, detailed assessment of a candidate.

Compare the candidate's CV to the job requirements below.
Assign numeric scores (0–100) and provide detailed reasoning for:
- Education
- Experience
- Skills
- Job-specific Fit

Return only valid JSON.

JOB REQUIREMENTS:
{self.job_requirements[:10000]}

CANDIDATE CV:
{cv_text[:12000]}
"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an HR evaluation system. Return ONLY valid JSON.",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                max_tokens=9000,
            )

            raw_output = response.choices[0].message.content.strip()
            clean_json = self._clean_json(raw_output)
            data = json.loads(clean_json)

            return CandidateAssessment(
                candidate_name=data.get("candidate_name", filename),
                filename=filename,
                overall_score=data.get("summary", {}).get("overall_fit_score", 0),
                fit_level=data.get("summary", {}).get("fit_level", "Unknown"),
                education_details=data.get("scoring_breakdown", {}).get("education", {}),
                experience_details=data.get("scoring_breakdown", {}).get("experience", {}),
                skills_details=data.get("scoring_breakdown", {}).get("skills", {}),
                job_fit_details=data.get("scoring_breakdown", {}).get("job_specific_fit", {}),
                weighted_score_total=data.get("weighted_score_total", 0),
                executive_summary=data.get("executive_summary", {}),
                recommendation=data.get("recommendation", {}),
                interview_focus_areas=data.get("interview_focus_areas", []),
                red_flags=data.get("red_flags", []),
                potential_concerns=data.get("potential_concerns", []),
                assessed_at=datetime.now().isoformat(),
            )

        except Exception as e:
            print(f"❌ Error in structured assessment: {e}")
            return CandidateAssessment(
                candidate_name="Error",
                filename=filename,
                overall_score=0,
                fit_level="Error",
                education_details={},
                experience_details={},
                skills_details={},
                job_fit_details={},
                weighted_score_total=0,
                executive_summary={"recommendation": "Assessment failed"},
                recommendation={"verdict": "Error", "rationale": str(e)},
                interview_focus_areas=[],
                red_flags=[],
                potential_concerns=[],
                assessed_at=datetime.now().isoformat(),
            )

    # ------------------- CRITICAL + TAILORING MODE -------------------

    def _assess_candidate_critical(self, filename: str, cv_text: str) -> Dict[str, Any]:
        """Critical narrative evaluation with adaptive criteria and weighting."""

        # ---------- 1️⃣ Semantic donor detection ----------
        donor_query = f"""
        Identify the main funding organization mentioned or implied in this tender.
        Return ONLY the donor name (e.g., 'World Bank', 'European Union', 'ADB', 'USAID', 'UNDP', 'AfDB', 'Unknown').
        If uncertain, answer exactly 'Unknown'.
        TENDER TEXT (excerpt):
        {self.job_requirements[:8000]}
        """
        donor_match = "Unknown"
        try:
            resp = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": donor_query}],
                temperature=0,
                max_tokens=10,
            )
            donor_match = resp.choices[0].message.content.strip()
        except Exception as e:
            print(f"⚠️ Donor semantic detection failed, fallback triggered: {e}")

        # ---------- 2️⃣ Regex fallback ----------
        text_lower = self.job_requirements.lower()
        donors = {
            "World Bank": r"\b(world\s*bank|wbg|ifc|ida|ibrd)\b",
            "European Union": r"\b(european\s*union|eu\s+delegation|europeaid|neighbourhood|dg\s*intl)\b",
            "Asian Development Bank": r"\b(asian\s+development\s+bank|adb)\b",
            "USAID": r"\b(usaid|united\s+states\s+agency\s+for\s+international\s+development)\b",
            "African Development Bank": r"\b(african\s+development\s+bank|afdb)\b",
            "UNDP": r"\b(undp|united\s+nations\s+development\s+programme)\b",
        }
        if donor_match == "Unknown":
            for name, pattern in donors.items():
                if re.search(pattern, text_lower):
                    donor_match = name
                    break
        if donor_match == "Unknown":
            donor_match = "General donor context"

        # ---------- 3️⃣ Build criteria table (custom or default) ----------
        if self.custom_criteria and 'criteria' in self.custom_criteria:
            # Build custom criteria table with weights
            criteria_rows = []
            criteria_rows.append("| **General Tender Context (20% weight)** | Understanding of project/tender context |  |  |  |")
            criteria_rows.append("|  | Familiarity with region/country context |  |  |  |")
            
            # Add custom criteria with their specific weights
            for criterion in self.custom_criteria['criteria']:
                weight_pct = criterion['weight']
                name = criterion['name']
                criteria_rows.append(f"| **{name} ({weight_pct}% weight)** | {name} |  |  |  |")
            
            criteria_table = "\n".join(criteria_rows)
            
            # Build weight explanation
            weights_explanation = "\n".join([
                f"- {c['name']}: {c['weight']}% (Rationale: {c['rationale']})"
                for c in self.custom_criteria['criteria']
            ])
            
            weight_instruction = f"""
IMPORTANT - CUSTOM WEIGHTING:
- General Tender Context: 20% total weight (split: 10% project understanding + 10% regional familiarity)
{weights_explanation}

When calculating the final weighted score:
1. Score each criterion from 0 to 1
2. Multiply each score by its weight percentage
3. Sum all weighted scores to get final score (should be 0-1.00)
"""
        else:
            # Default criteria table (equal weights within 80%)
            criteria_table = f"""| **General Tender Context (20% weight)** | Understanding of project/tender context |  |  |  |
|  | Familiarity with region/country context |  |  |  |
| **Specific Expert Requirements (80% weight)** | Team leadership and management |  |  |  |
|  | Relevant domain expertise |  |  |  |
|  | Technical or regulatory knowledge |  |  |  |
|  | Donor project experience ({donor_match}) |  |  |  |
|  | Communication and coordination skills |  |  |  |
|  | Educational background |  |  |  |
|  | Analytical and reporting skills |  |  |  |
|  | Language proficiency |  |  |  |"""
            weight_instruction = "Calculate weighted score: 20% for general context, 80% divided equally among expert requirements."

        # ---------- 4️⃣ Build the prompt ----------
        prompt = f"""
You are a senior evaluator assessing candidates for a tender funded by **{donor_match}**.

Perform a detailed, evidence-based critical evaluation of the candidate's CV
against the JOB REQUIREMENTS and contextualize every criterion according to {donor_match}'s
typical focus and terminology.
Do not mention other donors (EU, ADB, etc.) unless explicitly stated in the tender or CV.
Focus exclusively on {donor_match} as the donor context.

---

### REQUIRED OUTPUT STRUCTURE (Markdown)

#### 🧭 Critical Evaluation – {filename}

**Evaluation Table**

| Section | Criteria | Score (0–1) | Confidence | Evaluator Commentary |
|----------|-----------|-------------|-------------|----------------------|
{criteria_table}

Each cell must contain a numeric score (0–1), confidence (High/Medium/Low), and concise commentary (1–3 sentences).

{weight_instruction}

After the table, show:
**Final Weighted Score: X.XX / 1.00**

---

**📊 Critical Summary**
Summarize alignment with {donor_match} project style and focus.

**📉 Evaluator Summary**
List 3–5 key takeaways (strengths, weaknesses, donor-fit).

**📌 Strengths & Weaknesses**
Provide at least 3 evidence-based strengths and 3 weaknesses.

**✂️ Tailoring Suggestions (How to Strengthen CV for This Role)**

**a. Rewriting & Emphasis Suggestions**
Suggest concrete text-level improvements to highlight donor-specific relevance.

**b. 🪶 Word Recommendations (Tender Keyword Alignment)**
Provide 5–10 recommended keyword alignments using {donor_match} phrasing.

| Current CV Wording | Recommended {donor_match} Keyword | Why |
|--------------------|----------------------------------|-----|
| "Project manager" | "Implementation support consultant under {donor_match} framework" | Aligns with {donor_match} terminology. |
| "Drafted reports" | "Prepared deliverables per {donor_match} quality control protocols" | Matches donor QA phrasing. |
| "Policy development" | "Institutional reform and capacity-building support" | Fits {donor_match} operational tone. |

---

### CONTEXT INPUTS

**Tender Context (20%)**
{self.job_requirements[:4000]}

**Candidate CV**
{cv_text[:9000]}
"""

        # ---------- 5️⃣ Call GPT ----------
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a senior evaluator for international tenders. "
                            "Produce a full markdown report with scores, commentary, and rewording recommendations. "
                            "Always include the evaluation table and numeric final score."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.15,
                max_tokens=8500,
            )

            report_text = response.choices[0].message.content.strip()
            match = re.search(r"Final Weighted Score.*?([0-9]\.[0-9]+)", report_text)
            final_score = float(match.group(1)) if match else 0.0

            return {
                "candidate_name": filename,
                "report": f"**Detected Donor Context:** {donor_match}\n\n" + report_text,
                "final_score": final_score,
            }

        except Exception as e:
            return {
                "candidate_name": filename,
                "report": f"❌ Error generating critical evaluation: {e}",
                "final_score": 0.0,
            }

    # ------------------- JSON CLEANER -------------------

    def _clean_json(self, content: str) -> str:
        """Extract clean JSON from model output."""
        content = content.strip()
        content = re.sub(r"^```(json)?", "", content)
        content = re.sub(r"```$", "", content)
        match = re.search(r"(\{.*\})", content, re.DOTALL)
        if match:
            return match.group(1).strip()
        return content
